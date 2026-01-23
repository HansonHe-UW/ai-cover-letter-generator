import io
from docx import Document
from fpdf import FPDF

def create_docx(data):
    """
    Creates a Word document from the given data.
    
    Args:
        data (dict): Contains 'body' (str), 'user_info' (dict), 'hr_info' (dict).
        
    Returns:
        io.BytesIO: Buffer containing the .docx file.
    """
    text = data.get('body', '')
    doc = Document()
    # Handle simple parsing: strict markdown to plain text conversion is complex,
    # so we will treat it as mostly plain text, but maybe handle newlines properly.
    
    # Add a title (optional, or just start)
    # doc.add_heading('Cover Letter', 0)
    
    # Split by newlines to handle paragraphs
    for para in text.split('\n'):
        if para.strip():
            doc.add_paragraph(para)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(data):
    """
    Creates a PDF document from the given data.
    
    Args:
        data (dict): Contains 'body' (str), 'user_info' (dict), 'hr_info' (dict).
        
    Returns:
        io.BytesIO: Buffer containing the .pdf file.
    """
    text = data.get('body', '')
    pdf = FPDF(format='A4')
    pdf.add_page()
    
    # Use a standard font that supports basic English
    # Note: FPDF standard fonts (Times, Arial, Courier) don't support Chinese characters well without loading a custom font.
    # The requirement said "Ensure it handles standard English fonts".
    pdf.set_font("Times", size=12)
    
    # Calculate effective page width
    effective_page_width = pdf.w - 2 * pdf.l_margin
    
    # Basic Markdown stripping or handling
    clean_text = text.replace('**', '').replace('__', '')
    
    # Text Normalization for Standard Fonts (Latin-1)
    replacements = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '--', # Em dash
        '\u2026': '...',# Ellipsis
    }
    for char, replacement in replacements.items():
        clean_text = clean_text.replace(char, replacement)
    
    # Final safety net: Force encoding to latin-1, replacing unknown chars with ?
    clean_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.multi_cell(effective_page_width, 8, clean_text)
    
    buffer = io.BytesIO()
    # output(dest='S') returns the byte string. We need to wrap it in BytesIO.
    # Note: In fpdf2, output() without name returns the bytes.
    pdf_bytes = pdf.output() 
    buffer.write(pdf_bytes)
    buffer.seek(0)
    return buffer

def create_latex(data):
    """
    Creates a LaTeX source file and returns the source code.
    
    Args:
        data (dict): Contains 'body' (str), 'user_info' (dict), 'hr_info' (dict).
        
    Returns:
        tuple: (io.BytesIO, str) -> (buffer, latex_source_code)
    """
    text = data.get('body', '')
    user_info = data.get('user_info', {})
    date_str = data.get('date_str', r'\today')
    
    # Simple LaTeX Template
    latex_template = r"""
\documentclass[11pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage{geometry}
\geometry{a4paper, margin=1in}
\usepackage{parskip}

\begin{document}

\textbf{%s} \\
%s | %s | %s \\
%s

\vspace{1cm}

%s

\vspace{0.5cm}

%s

\end{document}
"""
    # Escape LaTeX special chars in body
    # This is a basic escaping, production needs more robust handling
    def latex_escape(s):
        chars = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\textasciicircum{}',
            '\\': r'\textbackslash{}',
        }
        return ''.join(chars.get(c, c) for c in s)

    escaped_body = latex_escape(text)
    
    # Fill template
    # Note: user_info and body might need escaping too, doing basic fill here
    # In a real app we'd carefully escape user_info values too
    latex_code = latex_template % (
        user_info.get('name', 'Name'),
        user_info.get('address', 'Address'),
        user_info.get('email', 'Email'),
        user_info.get('phone', 'Phone'),
        user_info.get('linkedin', 'LinkedIn'),
        date_str,
        escaped_body
    )
    
    buffer = io.BytesIO()
    buffer.write(latex_code.encode('utf-8'))
    buffer.seek(0)
    
    return buffer, latex_code
